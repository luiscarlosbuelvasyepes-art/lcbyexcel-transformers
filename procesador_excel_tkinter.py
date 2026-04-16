import math
import re
import tkinter as tk
from pathlib import Path
from tkinter import filedialog, messagebox, ttk

import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

try:
    from reportlab.lib.pagesizes import letter as rl_letter, A4 as rl_a4
    from reportlab.lib.units import cm as rl_cm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Spacer, PageBreak
    from reportlab.lib import colors as rl_colors
    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm as DocxCm
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False


PERIODOS = ["001", "002", "003", "FINAL"]
METRICAS = ["PROM", "PUESTO", "REPR", "ART", "CNT", "CPM", "EDF", "EMPIN"]
ALIAS_METRICAS = {
    "prom": "PROM",
    "promedio": "PROM",
    "puesto": "PUESTO",
    "repr": "REPR",
    "art": "ART",
    "cnt": "CNT",
    "cpm": "CPM",
    "edf": "EDF",
    "empi": "EMPIN",
    "empin": "EMPIN",
}
ALIAS_ESTUDIANTE = [
    "estudiante",
    "alumno",
    "alumna",
    "nombre",
    "nombres",
    "est",
    "id",
    "identificacion",
    "codigo",
    "documento",
]
ALIAS_PERIODO = ["periodo", "periodo academico", "periodo aca", "corte", "lapso"]
PATRONES_NO_ESTUDIANTE = [
    r"^\d+\.?\s*desempe[nñ]o",
    r"^desempe[nñ]o",
    r"consolidado",
    r"del curso",
    r"^grupo\b",
    r"totales?",
    r"resumen",
]
ESTUDIANTES_POR_FILA = 7
FILAS_POR_HOJA = 3
BLOQUE_ANCHO = 6
BLOQUE_ALTO = 12
TAMANIO_PAGINAS = {
    "carta": (21.59, 27.94),
    "a4":    (21.00, 29.70),
}


def parece_numero(valor: object) -> bool:
    if pd.isna(valor):
        return False
    texto = normalizar_texto(valor).replace(",", ".")
    if not texto:
        return False
    return bool(re.fullmatch(r"[-+]?\d+(\.\d+)?", texto))


def contiene_letras(valor: object) -> bool:
    texto = normalizar_texto(valor)
    return bool(re.search(r"[A-Za-zÁÉÍÓÚáéíóúÑñ]", texto))


def es_etiqueta_estudiante_valida(texto: str) -> bool:
    if not texto:
        return False

    clave = clave_texto(texto)
    if len(clave) < 4:
        return False

    if not contiene_letras(texto):
        return False

    for patron in PATRONES_NO_ESTUDIANTE:
        if re.search(patron, clave):
            return False

    return True


def es_columna_auxiliar(columna: object) -> bool:
    return bool(re.fullmatch(r"col_\d+(?:_\d+)?", clave_texto(columna)))


def normalizar_texto(valor: object) -> str:
    if pd.isna(valor):
        return ""
    texto = str(valor).strip()
    texto = re.sub(r"\s+", " ", texto)
    return texto


def clave_texto(valor: object) -> str:
    texto = normalizar_texto(valor).lower()
    texto = (
        texto.replace("á", "a")
        .replace("é", "e")
        .replace("í", "i")
        .replace("ó", "o")
        .replace("ú", "u")
    )
    return texto


def ruta_recurso(nombre: str) -> Path:
    """Devuelve la ruta correcta del recurso tanto en desarrollo como en .exe."""
    base = getattr(__import__("sys"), "_MEIPASS", None)
    if base:
        return Path(base) / nombre
    return Path(nombre)


def nombre_columna_unico(nombre: str, usados: set[str]) -> str:
    base = nombre or "COL"
    candidato = base
    sufijo = 2
    while candidato in usados:
        candidato = f"{base}_{sufijo}"
        sufijo += 1
    usados.add(candidato)
    return candidato


class GeneradorBoletinesApp:
    def __init__(self, root: tk.Tk) -> None:
        try:
            self.root = root
            self.root.title("GENERADOR DE INFORMES PERSONALIZADOS")

            # Ajustar tamaño y mínimo según la resolución real de la pantalla.
            ancho_pantalla = self.root.winfo_screenwidth()
            alto_pantalla  = self.root.winfo_screenheight()
            ancho_ventana  = max(860, min(int(ancho_pantalla * 0.88), 1600))
            alto_ventana   = max(560, min(int(alto_pantalla  * 0.88), 1000))
            x_pos = (ancho_pantalla  - ancho_ventana) // 2
            y_pos = max(0, (alto_pantalla - alto_ventana) // 2)
            self.root.geometry(f"{ancho_ventana}x{alto_ventana}+{x_pos}+{y_pos}")
            self.root.minsize(740, 500)
            self.root.resizable(True, True)
            self.root.configure(bg="#f4f7fb")

            # Factor de escala para fuentes basado en el ancho de la pantalla.
            self._escala = max(0.75, min(ancho_pantalla / 1920, 1.5))

            self.archivo_origen: Path | None = None
            self.df_raw: pd.DataFrame | None = None
            self.df_datos: pd.DataFrame | None = None
            self.reportes: dict[str, pd.DataFrame] = {}
            self.metricas_reporte: list[str] = METRICAS.copy()
            self.estudiantes_por_fila = tk.IntVar(value=ESTUDIANTES_POR_FILA)
            self.filas_por_hoja = tk.IntVar(value=FILAS_POR_HOJA)
            self.formato_salida = tk.StringVar(value="excel")
            self.tamano_pagina = tk.StringVar(value="carta")
            self.orientacion_pagina = tk.StringVar(value="horizontal")
            self.pagina_ancho_cm = tk.StringVar(value="21.59")
            self.pagina_alto_cm = tk.StringVar(value="27.94")
            self.periodos_vars: dict[str, tk.BooleanVar] = {
                "001":   tk.BooleanVar(value=True),
                "002":   tk.BooleanVar(value=False),
                "003":   tk.BooleanVar(value=False),
                "FINAL": tk.BooleanVar(value=False),
            }
            self.periodos_activos: list[str] = ["001"]

            self._configurar_estilos()
            self._crear_interfaz()
        except Exception as error:
            messagebox.showerror("Error", f"No se pudo inicializar la aplicacion:\n{error}")
            raise

    def _configurar_estilos(self) -> None:
        try:
            e = self._escala
            tam_titulo  = max(12, int(18 * e))
            tam_normal  = max(8,  int(10 * e))
            tam_boton   = max(8,  int(10 * e))
            tam_fila    = max(18, int(24 * e))
            pad_boton   = max(4,  int(8  * e))

            estilo = ttk.Style()
            estilo.theme_use("clam")
            estilo.configure("App.TFrame", background="#f4f7fb")
            estilo.configure(
                "Title.TLabel",
                background="#f4f7fb",
                foreground="#1f2937",
                font=("Segoe UI", tam_titulo, "bold"),
            )
            estilo.configure(
                "Status.TLabel",
                background="#f4f7fb",
                foreground="#334155",
                font=("Segoe UI", tam_normal),
            )
            estilo.configure(
                "Accent.TButton",
                font=("Segoe UI", tam_boton, "bold"),
                padding=pad_boton,
            )
            estilo.configure(
                "Treeview",
                background="#ffffff",
                fieldbackground="#ffffff",
                foreground="#111827",
                rowheight=tam_fila,
            )
            estilo.configure("Treeview.Heading", font=("Segoe UI", tam_normal, "bold"))
        except Exception as error:
            messagebox.showerror("Error", f"No se pudieron aplicar estilos:\n{error}")
            raise

    def _crear_interfaz(self) -> None:
        try:
            # Marco principal con grid para que la tabla se expanda en cualquier pantalla.
            self.root.columnconfigure(0, weight=1)
            self.root.rowconfigure(0, weight=1)
            marco = ttk.Frame(self.root, style="App.TFrame", padding=int(16 * self._escala))
            marco.grid(row=0, column=0, sticky="nsew")
            marco.columnconfigure(0, weight=1)
            marco.rowconfigure(99, weight=1)  # fila del contenedor de tabla

            # Conserva referencias para evitar que las imagenes sean recolectadas por GC.
            self.logo_institucional_img: tk.PhotoImage | None = None
            self.logo_personal_img: tk.PhotoImage | None = None

            encabezado = ttk.Frame(marco, style="App.TFrame")
            encabezado.pack(fill="x", pady=(0, 8))

            lbl_logo_izq = ttk.Label(encabezado, text="", style="Status.TLabel")
            lbl_logo_izq.grid(row=0, column=0, sticky="w")

            ttk.Label(
                encabezado,
                text="GENERADOR DE INFORMES PERSONALIZADOS",
                style="Title.TLabel",
            ).grid(row=0, column=1, sticky="n")

            lbl_logo_der = ttk.Label(encabezado, text="", style="Status.TLabel")
            lbl_logo_der.grid(row=0, column=2, sticky="e")

            encabezado.columnconfigure(1, weight=1)

            ruta_logo_institucional = ruta_recurso("LOGO ACTUALIZADO.png")
            if ruta_logo_institucional.exists():
                self.logo_institucional_img = tk.PhotoImage(file=str(ruta_logo_institucional))
                _s = max(3, int(5 / self._escala))
                self.logo_institucional_img = self.logo_institucional_img.subsample(_s, _s)
                lbl_logo_izq.configure(image=self.logo_institucional_img)
            else:
                lbl_logo_izq.configure(text="LOGO INSTITUCIONAL")

            ruta_logo_personal = ruta_recurso("LOGO LUISCARLOS.png")
            if ruta_logo_personal.exists():
                self.logo_personal_img = tk.PhotoImage(file=str(ruta_logo_personal))
                _s2 = max(2, int(3 / self._escala))
                self.logo_personal_img = self.logo_personal_img.subsample(_s2, _s2)
                lbl_logo_der.configure(image=self.logo_personal_img)
            else:
                lbl_logo_der.configure(text="LCBY")

            self.lbl_estado = ttk.Label(
                marco,
                text="Estado: Ningun archivo cargado.",
                style="Status.TLabel",
            )
            self.lbl_estado.pack(anchor="w", pady=(0, 10))

            botones = ttk.Frame(marco, style="App.TFrame")
            botones.pack(fill="x", pady=(0, 12))

            ttk.Button(
                botones,
                text="Cargar Excel",
                style="Accent.TButton",
                command=self.cargar_excel,
            ).grid(row=0, column=0, padx=(0, 8), pady=4, sticky="w")

            ttk.Button(
                botones,
                text="Generar Reportes",
                style="Accent.TButton",
                command=self.generar_reportes,
            ).grid(row=0, column=1, padx=8, pady=4, sticky="w")

            ttk.Button(
                botones,
                text="Guardar Archivo",
                style="Accent.TButton",
                command=self.guardar_archivo,
            ).grid(row=0, column=2, padx=8, pady=4, sticky="w")

            config = ttk.Frame(marco, style="App.TFrame")
            config.pack(anchor="w", pady=(0, 10))

            ttk.Label(config, text="Estudiantes por fila:", style="Status.TLabel").grid(
                row=0, column=0, sticky="w", padx=(0, 6)
            )
            ttk.Spinbox(
                config,
                from_=1,
                to=12,
                width=5,
                textvariable=self.estudiantes_por_fila,
            ).grid(row=0, column=1, sticky="w", padx=(0, 12))

            ttk.Label(config, text="Filas por hoja:", style="Status.TLabel").grid(
                row=0, column=2, sticky="w", padx=(0, 6)
            )
            ttk.Spinbox(
                config,
                from_=1,
                to=10,
                width=5,
                textvariable=self.filas_por_hoja,
            ).grid(row=0, column=3, sticky="w")

            frm_formato = ttk.Frame(marco, style="App.TFrame")
            frm_formato.pack(anchor="w", pady=(0, 6))
            ttk.Label(frm_formato, text="Formato de salida:", style="Status.TLabel").grid(
                row=0, column=0, sticky="w", padx=(0, 8)
            )
            for _col, (_etiq, _val) in enumerate(
                [("Excel", "excel"), ("PDF", "pdf"), ("Word", "word")], start=1
            ):
                ttk.Radiobutton(
                    frm_formato,
                    text=_etiq,
                    variable=self.formato_salida,
                    value=_val,
                ).grid(row=0, column=_col, padx=6, sticky="w")

            frm_pagina = ttk.Frame(marco, style="App.TFrame")
            frm_pagina.pack(anchor="w", pady=(0, 4))
            ttk.Label(frm_pagina, text="Tamaño de página:", style="Status.TLabel").grid(
                row=0, column=0, sticky="w", padx=(0, 8)
            )
            for _col, (_etiq, _val) in enumerate(
                [("Carta", "carta"), ("A4", "a4"), ("Personalizado", "custom")], start=1
            ):
                ttk.Radiobutton(
                    frm_pagina, text=_etiq,
                    variable=self.tamano_pagina, value=_val,
                    command=self._toggle_tamano_custom,
                ).grid(row=0, column=_col, padx=6, sticky="w")

            frm_orientacion = ttk.Frame(marco, style="App.TFrame")
            frm_orientacion.pack(anchor="w", pady=(0, 6))
            ttk.Label(frm_orientacion, text="Orientacion:", style="Status.TLabel").grid(
                row=0, column=0, sticky="w", padx=(0, 8)
            )
            for _col, (_etiq, _val) in enumerate(
                [("Vertical", "vertical"), ("Horizontal", "horizontal")], start=1
            ):
                ttk.Radiobutton(
                    frm_orientacion,
                    text=_etiq,
                    variable=self.orientacion_pagina,
                    value=_val,
                ).grid(row=0, column=_col, padx=6, sticky="w")

            self.frm_custom = ttk.Frame(marco, style="App.TFrame")
            ttk.Label(self.frm_custom, text="Ancho (cm):", style="Status.TLabel").grid(
                row=0, column=0, sticky="w", padx=(0, 4)
            )
            ttk.Entry(self.frm_custom, textvariable=self.pagina_ancho_cm, width=8).grid(
                row=0, column=1, padx=(0, 14)
            )
            ttk.Label(self.frm_custom, text="Alto (cm):", style="Status.TLabel").grid(
                row=0, column=2, sticky="w", padx=(0, 4)
            )
            ttk.Entry(self.frm_custom, textvariable=self.pagina_alto_cm, width=8).grid(
                row=0, column=3
            )

            frm_periodos = ttk.Frame(marco, style="App.TFrame")
            frm_periodos.pack(anchor="w", pady=(4, 4))
            ttk.Label(frm_periodos, text="Periodos a incluir:", style="Status.TLabel").grid(
                row=0, column=0, sticky="w", padx=(0, 8)
            )
            for _col, _periodo in enumerate(PERIODOS, start=1):
                ttk.Checkbutton(
                    frm_periodos,
                    text=_periodo,
                    variable=self.periodos_vars[_periodo],
                ).grid(row=0, column=_col, padx=6, sticky="w")

            ttk.Button(
                frm_periodos,
                text="Seleccionar todos",
                command=self._seleccionar_todos_periodos,
            ).grid(row=0, column=len(PERIODOS) + 1, padx=(12, 6), sticky="w")

            ttk.Button(
                frm_periodos,
                text="Limpiar selección",
                command=self._limpiar_seleccion_periodos,
            ).grid(row=0, column=len(PERIODOS) + 2, padx=6, sticky="w")

            self.lbl_resumen = ttk.Label(
                marco,
                text="Resumen: sin datos procesados.",
                style="Status.TLabel",
            )
            self.lbl_resumen.pack(anchor="w", pady=(0, 6))

            # Contenedor de tabla: usa grid directamente en marco para expandirse.
            contenedor_tabla = ttk.Frame(marco, style="App.TFrame")
            contenedor_tabla.pack(fill="both", expand=True)
            contenedor_tabla.rowconfigure(0, weight=1)
            contenedor_tabla.columnconfigure(0, weight=1)

            self.tabla = ttk.Treeview(contenedor_tabla, show="headings")
            self.tabla.grid(row=0, column=0, sticky="nsew")

            barra_y = ttk.Scrollbar(contenedor_tabla, orient="vertical", command=self.tabla.yview)
            barra_y.grid(row=0, column=1, sticky="ns")

            barra_x = ttk.Scrollbar(contenedor_tabla, orient="horizontal", command=self.tabla.xview)
            barra_x.grid(row=1, column=0, sticky="ew")

            self.tabla.configure(yscrollcommand=barra_y.set, xscrollcommand=barra_x.set)

            pie = ttk.Frame(marco, style="App.TFrame")
            pie.pack(fill="x", pady=(8, 0))
            ttk.Label(
                pie,
                text="Luis Carlos Buelvas Yepes",
                style="Status.TLabel",
            ).pack(side="left")
            ttk.Label(
                pie,
                text="Soluciones Educativas YAIMAR",
                style="Status.TLabel",
            ).pack(side="right")

            self._toggle_tamano_custom()
        except Exception as error:
            messagebox.showerror("Error", f"No se pudo construir la interfaz:\n{error}")
            raise

    def _toggle_tamano_custom(self) -> None:
        if self.tamano_pagina.get() == "custom":
            self.frm_custom.pack(anchor="w", pady=(0, 6))
        else:
            self.frm_custom.pack_forget()

    def _seleccionar_todos_periodos(self) -> None:
        for periodo in PERIODOS:
            self.periodos_vars[periodo].set(True)

    def _limpiar_seleccion_periodos(self) -> None:
        for periodo in PERIODOS:
            self.periodos_vars[periodo].set(False)

    def _obtener_periodos_seleccionados(self) -> list[str]:
        seleccionados = [p for p in PERIODOS if self.periodos_vars[p].get()]
        if not seleccionados:
            raise ValueError(
                "Debes seleccionar al menos un periodo para generar los reportes."
            )
        return seleccionados

    def _obtener_tamano_pagina(self) -> tuple[float, float]:
        modo = self.tamano_pagina.get()
        if modo in TAMANIO_PAGINAS:
            return TAMANIO_PAGINAS[modo]
        try:
            ancho = float(self.pagina_ancho_cm.get())
            alto  = float(self.pagina_alto_cm.get())
            if ancho <= 0 or alto <= 0:
                raise ValueError("Dimensiones deben ser positivas.")
            return ancho, alto
        except ValueError as error:
            raise ValueError(f"Dimensiones de pagina invalidas: {error}") from error

    def _es_orientacion_horizontal(self) -> bool:
        return self.orientacion_pagina.get() == "horizontal"

    def _obtener_tamano_pagina_orientado(self) -> tuple[float, float]:
        ancho_cm, alto_cm = self._obtener_tamano_pagina()
        if self._es_orientacion_horizontal():
            return max(ancho_cm, alto_cm), min(ancho_cm, alto_cm)
        return min(ancho_cm, alto_cm), max(ancho_cm, alto_cm)

    def _actualizar_estado(self, mensaje: str) -> None:
        self.lbl_estado.config(text=mensaje)

    def _actualizar_resumen(self, mensaje: str) -> None:
        self.lbl_resumen.config(text=mensaje)

    def _mostrar_dataframe(self, df: pd.DataFrame) -> None:
        self.tabla.delete(*self.tabla.get_children())
        columnas = [str(columna) for columna in df.columns]
        self.tabla["columns"] = columnas

        for indice, columna in enumerate(columnas):
            self.tabla.heading(columna, text=columna)
            self.tabla.column(columna, anchor="center", width=self._calcular_ancho(df, indice, columna))

        for fila in df.fillna("").itertuples(index=False):
            self.tabla.insert("", "end", values=list(fila))

    def _calcular_ancho(self, df: pd.DataFrame, indice_columna: int, encabezado: str) -> int:
        longitud_encabezado = len(str(encabezado))
        if df.empty or indice_columna >= df.shape[1]:
            longitud_datos = 0
        else:
            serie = df.iloc[:, indice_columna].apply(lambda valor: "" if pd.isna(valor) else str(valor))
            longitud_datos = serie.map(len).max()
        return min(max(90, max(longitud_encabezado, int(longitud_datos or 0)) * 8 + 16), 260)

    def cargar_excel(self) -> None:
        try:
            ruta = filedialog.askopenfilename(
                title="Seleccionar archivo Excel",
                filetypes=[("Archivos Excel", "*.xlsx")],
            )
            if not ruta:
                self._actualizar_estado("Estado: Carga cancelada por el usuario.")
                return

            self.archivo_origen = Path(ruta)
            self.df_raw = pd.read_excel(ruta, header=None, dtype=object)
            self.df_datos = self._preparar_dataframe(self.df_raw)
            self.reportes = {}

            self._mostrar_dataframe(self.df_datos.head(50))
            self._actualizar_estado(f"Estado: Archivo cargado -> {self.archivo_origen.name}")
            self._actualizar_resumen(
                f"Resumen: {len(self.df_datos)} filas detectadas y {len(self.df_datos.columns)} columnas utiles."
            )
            messagebox.showinfo("Exito", "Archivo cargado correctamente.")
        except Exception as error:
            self.archivo_origen = None
            self.df_raw = None
            self.df_datos = None
            self.reportes = {}
            self._actualizar_estado("Estado: Error al cargar el archivo.")
            messagebox.showerror("Error", f"No se pudo cargar el archivo:\n{error}")

    def _preparar_dataframe(self, df_raw: pd.DataFrame) -> pd.DataFrame:
        df_limpio = df_raw.dropna(how="all").dropna(axis=1, how="all").reset_index(drop=True)
        if df_limpio.empty:
            raise ValueError("El archivo no contiene datos utilizables.")

        fila_encabezado = self._detectar_fila_encabezado(df_limpio)
        fila = df_limpio.iloc[fila_encabezado].tolist()
        usados: set[str] = set()
        encabezados: list[str] = []

        for indice, valor in enumerate(fila, start=1):
            texto = normalizar_texto(valor) or f"COL_{indice}"
            encabezados.append(nombre_columna_unico(texto, usados))

        df_datos = df_limpio.iloc[fila_encabezado + 1 :].copy()
        df_datos.columns = encabezados
        df_datos = df_datos.dropna(how="all").reset_index(drop=True)

        if df_datos.empty:
            raise ValueError("No se encontraron filas de datos debajo del encabezado detectado.")

        return df_datos

    def _detectar_fila_encabezado(self, df: pd.DataFrame) -> int:
        mejor_indice = 0
        mejor_puntaje = -1
        limite = min(len(df), 20)

        for indice in range(limite):
            puntaje = 0
            for valor in df.iloc[indice].tolist():
                clave = clave_texto(valor)
                if clave in ALIAS_METRICAS:
                    puntaje += 4
                elif clave in ALIAS_PERIODO:
                    puntaje += 3
                elif clave in ALIAS_ESTUDIANTE:
                    puntaje += 3
                elif normalizar_texto(valor).upper() in PERIODOS:
                    puntaje += 1
            if puntaje > mejor_puntaje:
                mejor_puntaje = puntaje
                mejor_indice = indice

        return mejor_indice

    def _detectar_columna_periodo(self, df: pd.DataFrame) -> str:
        mejor_columna = ""
        mejor_puntaje = -1.0

        for columna in df.columns:
            nombre = clave_texto(columna)
            valores = df[columna].apply(normalizar_texto)
            valores = valores[valores != ""]
            if valores.empty:
                continue

            coincidencias = valores.apply(lambda valor: valor.upper() in PERIODOS).sum()
            puntaje = coincidencias / len(valores)
            if nombre in ALIAS_PERIODO:
                puntaje += 1

            if puntaje > mejor_puntaje:
                mejor_puntaje = puntaje
                mejor_columna = columna

        if not mejor_columna or mejor_puntaje <= 0:
            raise ValueError("No fue posible identificar la columna del periodo.")
        return mejor_columna

    def _detectar_columnas_estudiante(self, df: pd.DataFrame, columna_periodo: str) -> tuple[str, str | None]:
        candidatos: list[tuple[float, str]] = []

        for columna in df.columns:
            if columna == columna_periodo:
                continue

            nombre = clave_texto(columna)
            serie = df[columna].apply(normalizar_texto)
            valores = serie[serie != ""]
            if valores.empty:
                continue

            longitud_media = valores.map(len).mean()
            variedad = valores.nunique() / max(len(valores), 1)
            puntaje = variedad + min(longitud_media / 20, 1)

            proporcion_letras = valores.apply(contiene_letras).mean()
            proporcion_numerica = valores.apply(parece_numero).mean()
            puntaje += proporcion_letras * 2
            puntaje -= proporcion_numerica * 2

            if any(alias in nombre for alias in ALIAS_ESTUDIANTE):
                puntaje += 3
            if valores.apply(lambda valor: valor.upper() in PERIODOS).mean() > 0.5:
                puntaje = -1

            candidatos.append((puntaje, columna))

        candidatos.sort(reverse=True)
        if not candidatos or candidatos[0][0] <= 0:
            raise ValueError("No fue posible identificar la columna del estudiante.")

        principal = candidatos[0][1]
        secundario = None
        for puntaje, columna in candidatos[1:]:
            nombre = clave_texto(columna)
            if puntaje > 0.5 and any(alias in nombre for alias in ["id", "codigo", "documento"]):
                secundario = columna
                break
        return principal, secundario

    def _detectar_columnas_metricas(self, df: pd.DataFrame, ignoradas: set[str]) -> dict[str, str]:
        mapeo: dict[str, str] = {}
        for columna in df.columns:
            if columna in ignoradas:
                continue

            clave = clave_texto(columna)
            if clave in ALIAS_METRICAS:
                mapeo[ALIAS_METRICAS[clave]] = columna
                continue

            serie = df[columna].apply(normalizar_texto)
            no_vacios = serie[serie != ""]
            if no_vacios.empty:
                continue

            proporcion_numerica = no_vacios.apply(parece_numero).mean()
            if proporcion_numerica >= 0.35:
                nombre_visible = normalizar_texto(columna).upper()
                if nombre_visible and not clave.startswith("col_"):
                    mapeo[nombre_visible] = columna
        return mapeo

    def _construir_etiqueta_estudiante(
        self,
        fila: pd.Series,
        columna_principal: str,
        columna_secundaria: str | None,
    ) -> str:
        principal = normalizar_texto(fila.get(columna_principal, ""))
        secundario = normalizar_texto(fila.get(columna_secundaria, "")) if columna_secundaria else ""
        if principal and secundario and principal != secundario:
            return f"{principal} ({secundario})"
        return principal or secundario

    def _es_vacio(self, valor: object) -> bool:
        return pd.isna(valor) or normalizar_texto(valor) == ""

    def _fusionar_columnas_auxiliares(self, df: pd.DataFrame, columna_periodo: str) -> pd.DataFrame:
        trabajo = df.copy()
        columnas = list(trabajo.columns)
        a_eliminar: list[str] = []

        periodos = trabajo[columna_periodo].apply(lambda valor: normalizar_texto(valor).upper())
        es_final = periodos == "FINAL"

        for indice, columna in enumerate(columnas):
            if indice == 0 or not es_columna_auxiliar(columna):
                continue

            destino = columnas[indice - 1]
            if es_columna_auxiliar(destino):
                continue

            destino_vacio = trabajo[destino].apply(self._es_vacio)
            auxiliar_con_dato = ~trabajo[columna].apply(self._es_vacio)

            # Prioriza completar filas del periodo FINAL con la columna auxiliar.
            mascara_final = es_final & destino_vacio & auxiliar_con_dato
            if mascara_final.any():
                trabajo.loc[mascara_final, destino] = trabajo.loc[mascara_final, columna]

            # Si aun hay vacios en otros periodos, tambien los completa.
            destino_vacio = trabajo[destino].apply(self._es_vacio)
            mascara_general = destino_vacio & auxiliar_con_dato
            if mascara_general.any():
                trabajo.loc[mascara_general, destino] = trabajo.loc[mascara_general, columna]

            a_eliminar.append(columna)

        if a_eliminar:
            trabajo = trabajo.drop(columns=a_eliminar, errors="ignore")

        return trabajo

    def _normalizar_tabla_fuente(self, df: pd.DataFrame) -> pd.DataFrame:
        columna_periodo = self._detectar_columna_periodo(df)
        df = self._fusionar_columnas_auxiliares(df, columna_periodo)

        columna_periodo = self._detectar_columna_periodo(df)
        columna_estudiante, columna_id = self._detectar_columnas_estudiante(df, columna_periodo)
        ignoradas = {columna_periodo, columna_estudiante}
        if columna_id:
            ignoradas.add(columna_id)

        columnas_metricas = self._detectar_columnas_metricas(df, ignoradas)
        if not columnas_metricas:
            raise ValueError("No se detectaron columnas de metricas en el archivo.")

        metricas_ordenadas: list[str] = []
        for metrica_base in METRICAS:
            if metrica_base in columnas_metricas:
                metricas_ordenadas.append(metrica_base)
        for metrica in columnas_metricas.keys():
            if metrica not in metricas_ordenadas:
                metricas_ordenadas.append(metrica)
        self.metricas_reporte = metricas_ordenadas

        trabajo = df.copy()
        trabajo[columna_estudiante] = trabajo[columna_estudiante].ffill()
        if columna_id:
            trabajo[columna_id] = trabajo[columna_id].ffill()

        trabajo["_periodo"] = trabajo[columna_periodo].apply(lambda valor: normalizar_texto(valor).upper())
        trabajo = trabajo[trabajo["_periodo"].isin(PERIODOS)].copy()
        if trabajo.empty:
            raise ValueError("No se encontraron filas con periodos 001, 002, 003 o FINAL.")

        trabajo["_estudiante"] = trabajo.apply(
            lambda fila: self._construir_etiqueta_estudiante(fila, columna_estudiante, columna_id),
            axis=1,
        )
        trabajo = trabajo[trabajo["_estudiante"].apply(es_etiqueta_estudiante_valida)].copy()
        if trabajo.empty:
            raise ValueError(
                "No se detectaron estudiantes validos. El archivo parece contener solo consolidados o resumenes."
            )

        columnas_finales = ["_estudiante", "_periodo"]
        renombrar: dict[str, str] = {}

        for metrica, columna_original in columnas_metricas.items():
            columnas_finales.append(columna_original)
            renombrar[columna_original] = metrica

        normalizado = trabajo[columnas_finales].rename(columns=renombrar).copy()

        for metrica in self.metricas_reporte:
            if metrica in normalizado.columns:
                normalizado[metrica] = pd.to_numeric(normalizado[metrica], errors="coerce")

        return normalizado

    def _crear_matriz_estudiante(self, df_estudiante: pd.DataFrame) -> pd.DataFrame:
        # dtype=object permite mezclar texto y valores numericos sin errores de tipo.
        periodos = self.periodos_activos
        matriz = pd.DataFrame(index=self.metricas_reporte, columns=periodos, dtype=object)
        matriz = matriz.fillna("")

        for _, fila in df_estudiante.iterrows():
            periodo = fila["_periodo"]
            if periodo not in periodos:
                continue
            for metrica in self.metricas_reporte:
                if metrica in df_estudiante.columns:
                    valor = fila.get(metrica)
                    if pd.notna(valor):
                        matriz.loc[metrica, periodo] = self._formatear_valor(valor)

        return matriz

    def _formatear_valor(self, valor: object) -> object:
        if pd.isna(valor):
            return ""
        if isinstance(valor, float):
            if math.isclose(valor, round(valor)):
                return int(round(valor))
            return round(valor, 2)
        return valor

    def _obtener_configuracion_pagina(self) -> tuple[int, int]:
        try:
            estudiantes = int(self.estudiantes_por_fila.get())
            filas = int(self.filas_por_hoja.get())
        except Exception as error:
            raise ValueError(f"Configuracion invalida de pagina: {error}") from error

        if estudiantes <= 0 or filas <= 0:
            raise ValueError("Estudiantes por fila y filas por hoja deben ser mayores a cero.")

        return estudiantes, filas

    def generar_reportes(self) -> None:
        try:
            if self.df_datos is None:
                messagebox.showwarning("Atencion", "Primero debes cargar un archivo Excel.")
                return

            self.periodos_activos = self._obtener_periodos_seleccionados()
            df_normalizado = self._normalizar_tabla_fuente(self.df_datos)
            reportes: dict[str, pd.DataFrame] = {}
            for estudiante, grupo in df_normalizado.groupby("_estudiante", sort=True):
                reportes[estudiante] = self._crear_matriz_estudiante(grupo)

            if not reportes:
                raise ValueError("No se generaron reportes individuales.")

            self.reportes = reportes
            vista = next(iter(reportes.values())).reset_index().rename(columns={"index": "Metrica"})
            self._mostrar_dataframe(vista)
            self._actualizar_estado("Estado: Reportes individuales generados correctamente.")
            self._actualizar_resumen(
                f"Resumen: {len(self.reportes)} estudiantes listos para exportar. Periodos: {', '.join(self.periodos_activos)}."
            )
            messagebox.showinfo("Exito", f"Se generaron {len(self.reportes)} reportes individuales.")
        except Exception as error:
            self.reportes = {}
            self._actualizar_estado("Estado: Error al generar los reportes.")
            messagebox.showerror("Error", f"No se pudieron generar los reportes:\n{error}")

    def guardar_archivo(self) -> None:
        try:
            if not self.reportes:
                messagebox.showwarning("Atencion", "Primero debes generar los reportes individuales.")
                return

            fmt = self.formato_salida.get()
            if fmt == "pdf" and not REPORTLAB_OK:
                messagebox.showerror(
                    "Dependencia faltante",
                    "Para exportar a PDF instala reportlab:\n  pip install reportlab",
                )
                return
            if fmt == "word" and not DOCX_OK:
                messagebox.showerror(
                    "Dependencia faltante",
                    "Para exportar a Word instala python-docx:\n  pip install python-docx",
                )
                return

            ext_map   = {"excel": ".xlsx", "pdf": ".pdf", "word": ".docx"}
            tipos_map = {
                "excel": [("Archivos Excel", "*.xlsx")],
                "pdf":   [("Archivos PDF",   "*.pdf")],
                "word":  [("Archivos Word",  "*.docx")],
            }
            ext  = ext_map[fmt]
            base = "boletines_estudiantes"
            if self.archivo_origen is not None:
                base = f"{self.archivo_origen.stem}_boletines"

            ruta = filedialog.asksaveasfilename(
                title="Guardar archivo de boletines",
                defaultextension=ext,
                initialfile=f"{base}{ext}",
                filetypes=tipos_map[fmt],
            )
            if not ruta:
                self._actualizar_estado("Estado: Guardado cancelado por el usuario.")
                return

            if fmt == "excel":
                self._crear_libro_excel().save(ruta)
            elif fmt == "pdf":
                self._crear_pdf(ruta)
            else:
                self._crear_word(ruta)

            self._actualizar_estado("Estado: Archivo final generado correctamente.")
            messagebox.showinfo("Exito", f"Archivo guardado en:\n{ruta}")
        except Exception as error:
            self._actualizar_estado("Estado: Error al guardar el archivo.")
            messagebox.showerror("Error", f"No se pudo guardar el archivo:\n{error}")

    def _crear_libro_excel(self) -> Workbook:
        libro = Workbook()
        libro.remove(libro.active)

        estudiantes_por_fila, filas_por_hoja = self._obtener_configuracion_pagina()
        periodos = self.periodos_activos
        columnas_por_bloque = 1 + len(periodos)

        # Alto exacto: 1 fila titulo + 1 fila encabezado periodos + 1 fila por metrica.
        # Se agrega 1 fila de separacion pequena entre filas de bloques (alto=8).
        alto_contenido   = 2 + len(self.metricas_reporte)
        alto_separacion  = 1
        alto_bloque      = alto_contenido + alto_separacion

        estudiantes = list(self.reportes.items())
        por_hoja = estudiantes_por_fila * filas_por_hoja

        for inicio in range(0, len(estudiantes), por_hoja):
            hoja = libro.create_sheet(title=f"Pagina_{inicio // por_hoja + 1}")
            self._configurar_hoja(hoja)

            for indice_bloque, (estudiante, matriz) in enumerate(estudiantes[inicio : inicio + por_hoja]):
                fila_bloque    = indice_bloque // estudiantes_por_fila
                columna_bloque = indice_bloque % estudiantes_por_fila
                fila_inicio    = 1 + fila_bloque * alto_bloque
                columna_inicio = 1 + columna_bloque * columnas_por_bloque
                self._escribir_bloque_estudiante(hoja, fila_inicio, columna_inicio, estudiante, matriz)

                # Ajustar la altura de cada fila del bloque y de la fila separadora.
                hoja.row_dimensions[fila_inicio].height     = 18  # titulo
                hoja.row_dimensions[fila_inicio + 1].height = 14  # encabezado periodos
                for _fo in range(2, alto_contenido):
                    hoja.row_dimensions[fila_inicio + _fo].height = 13  # metricas
                hoja.row_dimensions[fila_inicio + alto_contenido].height = 8  # separacion

            self._ajustar_anchos_hoja(hoja, columnas_por_bloque)

        return libro

    def _configurar_hoja(self, hoja) -> None:
        modo = self.tamano_pagina.get()
        if modo == "carta":
            hoja.page_setup.paperSize = hoja.PAPERSIZE_LETTER
        else:
            hoja.page_setup.paperSize = hoja.PAPERSIZE_A4
        hoja.page_setup.orientation = "landscape" if self._es_orientacion_horizontal() else "portrait"
        hoja.page_setup.fitToWidth = 1
        hoja.page_setup.fitToHeight = 1
        hoja.sheet_view.showGridLines = False
        hoja.page_margins.left = 0.2
        hoja.page_margins.right = 0.2
        hoja.page_margins.top = 0.3
        hoja.page_margins.bottom = 0.3

    def _escribir_bloque_estudiante(self, hoja, fila_inicio: int, columna_inicio: int, estudiante: str, matriz: pd.DataFrame) -> None:
        periodos = self.periodos_activos
        relleno_titulo = PatternFill("solid", fgColor="D9D9D9")
        relleno_encabezado = PatternFill("solid", fgColor="EDEDED")
        borde = Border(
            left=Side(style="thin", color="808080"),
            right=Side(style="thin", color="808080"),
            top=Side(style="thin", color="808080"),
            bottom=Side(style="thin", color="808080"),
        )
        centrado = Alignment(horizontal="center", vertical="center", wrap_text=True)
        fuente_titulo = Font(name="Calibri", size=10, bold=True)
        fuente_normal = Font(name="Calibri", size=9)
        ultima_columna = columna_inicio + len(periodos)

        hoja.merge_cells(
            start_row=fila_inicio,
            start_column=columna_inicio,
            end_row=fila_inicio,
            end_column=ultima_columna,
        )
        celda_titulo = hoja.cell(row=fila_inicio, column=columna_inicio, value=estudiante)
        celda_titulo.fill = relleno_titulo
        celda_titulo.border = borde
        celda_titulo.alignment = centrado
        celda_titulo.font = fuente_titulo

        for columna_offset, periodo in enumerate(["", *periodos]):
            celda = hoja.cell(row=fila_inicio + 1, column=columna_inicio + columna_offset, value=periodo)
            celda.fill = relleno_encabezado
            celda.border = borde
            celda.alignment = centrado
            celda.font = fuente_titulo

        for fila_offset, metrica in enumerate(self.metricas_reporte, start=2):
            celda_metrica = hoja.cell(row=fila_inicio + fila_offset, column=columna_inicio, value=metrica)
            celda_metrica.fill = relleno_encabezado
            celda_metrica.border = borde
            celda_metrica.alignment = centrado
            celda_metrica.font = fuente_titulo

            for columna_offset, periodo in enumerate(periodos, start=1):
                valor = matriz.loc[metrica, periodo] if periodo in matriz.columns else ""
                celda_valor = hoja.cell(
                    row=fila_inicio + fila_offset,
                    column=columna_inicio + columna_offset,
                    value=valor,
                )
                celda_valor.border = borde
                celda_valor.alignment = centrado
                celda_valor.font = fuente_normal

        for columna in range(columna_inicio, ultima_columna + 1):
            hoja.cell(row=fila_inicio, column=columna).border = borde

    def _ajustar_anchos_hoja(self, hoja, columnas_por_bloque: int) -> None:
        for columna in range(1, hoja.max_column + 1):
            letra = get_column_letter(columna)
            maximo = 0
            for celda in hoja[letra]:
                valor = "" if celda.value is None else str(celda.value)
                maximo = max(maximo, len(valor))

            # Primera columna del bloque: etiqueta de la metrica.
            if (columna - 1) % columnas_por_bloque == 0:
                hoja.column_dimensions[letra].width = min(max(maximo + 2, 10), 18)
            # Columnas de periodos seleccionados.
            else:
                hoja.column_dimensions[letra].width = min(max(maximo + 2, 7), 12)

    # ------------------------------------------------------------------ PDF --
    def _crear_pdf(self, ruta: str) -> None:
        ancho_cm, alto_cm = self._obtener_tamano_pagina_orientado()
        ancho_pt = ancho_cm * rl_cm
        alto_pt  = alto_cm * rl_cm
        estudiantes_por_fila, filas_por_hoja = self._obtener_configuracion_pagina()
        por_hoja = estudiantes_por_fila * filas_por_hoja
        metricas = self.metricas_reporte
        periodos = self.periodos_activos
        num_per  = len(periodos)

        margen_pt    = 0.5 * rl_cm
        disponible   = ancho_pt - 2 * margen_pt
        gap_pt       = 3
        bloque_pt    = (disponible - gap_pt * (estudiantes_por_fila - 1)) / estudiantes_por_fila
        col_label_pt = bloque_pt * 0.30
        col_per_pt   = (bloque_pt - col_label_pt) / num_per
        col_anchos   = [col_label_pt] + [col_per_pt] * num_per
        altos        = [14, 11] + [10] * len(metricas)

        estilo_inner = TableStyle([
            ("FONTNAME",   (0, 0), (-1, -1), "Helvetica"),
            ("FONTSIZE",   (0, 0), (-1, -1), 7),
            ("ALIGN",      (0, 0), (-1, -1), "CENTER"),
            ("VALIGN",     (0, 0), (-1, -1), "MIDDLE"),
            ("GRID",       (0, 0), (-1, -1), 0.4, rl_colors.grey),
            ("BACKGROUND", (0, 0), (-1, 0),  rl_colors.HexColor("#D9D9D9")),
            ("BACKGROUND", (0, 1), (-1, 1),  rl_colors.HexColor("#EDEDED")),
            ("BACKGROUND", (0, 2), (0, -1),  rl_colors.HexColor("#EDEDED")),
            ("FONTNAME",   (0, 0), (-1, 1),  "Helvetica-Bold"),
            ("FONTNAME",   (0, 2), (0, -1),  "Helvetica-Bold"),
            ("SPAN",       (0, 0), (-1, 0)),
        ])

        def _tabla_est(nombre, matriz):
            encab = [nombre] + [""] * num_per
            sub   = [""] + periodos
            datos = [encab, sub]
            for m in metricas:
                fila = [m]
                for p in periodos:
                    v = matriz.loc[m, p] if (m in matriz.index and p in matriz.columns) else ""
                    fila.append("" if v == "" else str(v))
                datos.append(fila)
            t = Table(datos, colWidths=col_anchos, rowHeights=altos)
            t.setStyle(estilo_inner)
            return t

        story = []
        todos = list(self.reportes.items())
        primera = True
        for inicio in range(0, len(todos), por_hoja):
            if not primera:
                story.append(PageBreak())
            primera = False
            grupo = todos[inicio:inicio + por_hoja]
            for fi in range(0, len(grupo), estudiantes_por_fila):
                fila_ests = grupo[fi:fi + estudiantes_por_fila]
                tablas = [_tabla_est(n, m) for n, m in fila_ests]
                while len(tablas) < estudiantes_por_fila:
                    n_r = 2 + len(metricas)
                    t_vacio = Table([[""] * (1 + num_per)] * n_r, colWidths=col_anchos, rowHeights=altos)
                    tablas.append(t_vacio)
                maestra = Table(
                    [tablas],
                    colWidths=[bloque_pt] * estudiantes_por_fila,
                    hAlign="LEFT",
                )
                maestra.setStyle(TableStyle([
                    ("LEFTPADDING",   (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING",  (0, 0), (-1, -1), gap_pt),
                    ("TOPPADDING",    (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                    ("VALIGN",        (0, 0), (-1, -1), "TOP"),
                ]))
                story.append(maestra)
                story.append(Spacer(1, 6))

        SimpleDocTemplate(
            ruta,
            pagesize=(ancho_pt, alto_pt),
            leftMargin=margen_pt, rightMargin=margen_pt,
            topMargin=margen_pt,  bottomMargin=margen_pt,
        ).build(story)

    # ----------------------------------------------------------------- Word --
    def _crear_word(self, ruta: str) -> None:
        ancho_cm, alto_cm = self._obtener_tamano_pagina_orientado()
        estudiantes_por_fila, filas_por_hoja = self._obtener_configuracion_pagina()
        por_hoja = estudiantes_por_fila * filas_por_hoja
        metricas = self.metricas_reporte
        periodos = self.periodos_activos
        cols_por_est = 1 + len(periodos)

        documento = DocxDocument()
        sec = documento.sections[0]
        sec.orientation = WD_ORIENT.LANDSCAPE if self._es_orientacion_horizontal() else WD_ORIENT.PORTRAIT
        sec.page_width    = DocxCm(ancho_cm)
        sec.page_height   = DocxCm(alto_cm)
        sec.left_margin   = DocxCm(0.8)
        sec.right_margin  = DocxCm(0.8)
        sec.top_margin    = DocxCm(0.8)
        sec.bottom_margin = DocxCm(0.8)

        todos = list(self.reportes.items())
        primera = True
        for inicio in range(0, len(todos), por_hoja):
            if not primera:
                documento.add_page_break()
            primera = False
            grupo = todos[inicio:inicio + por_hoja]
            num_filas_tabla = 2 + len(metricas)

            for fi in range(0, len(grupo), estudiantes_por_fila):
                fila_ests = grupo[fi:fi + estudiantes_por_fila]
                n = len(fila_ests)
                total_cols = n * cols_por_est

                tabla = documento.add_table(rows=num_filas_tabla, cols=total_cols)
                tabla.style = "Table Grid"

                ancho_disp = ancho_cm - 1.6
                ancho_lbl = (ancho_disp / n) * 0.30
                ancho_per = ((ancho_disp / n) - ancho_lbl) / max(len(periodos), 1)

                for col_est, (nombre_est, matriz) in enumerate(fila_ests):
                    cb = col_est * cols_por_est

                    celda_tit = tabla.cell(0, cb)
                    for c in range(1, cols_por_est):
                        celda_tit = celda_tit.merge(tabla.cell(0, cb + c))
                    p = celda_tit.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run(str(nombre_est).upper())
                    r.bold = True
                    r.font.size = Pt(7)

                    for off, txt in enumerate([""] + periodos):
                        p = tabla.cell(1, cb + off).paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r = p.add_run(txt)
                        r.bold = True
                        r.font.size = Pt(7)

                    for fo, metrica in enumerate(metricas, start=2):
                        p = tabla.cell(fo, cb).paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r = p.add_run(metrica)
                        r.bold = True
                        r.font.size = Pt(7)
                        for co, periodo in enumerate(periodos, start=1):
                            v = ""
                            if metrica in matriz.index and periodo in matriz.columns:
                                raw = matriz.loc[metrica, periodo]
                                v = "" if raw == "" else str(raw)
                            p2 = tabla.cell(fo, cb + co).paragraphs[0]
                            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            r2 = p2.add_run(v)
                            r2.font.size = Pt(7)

                for col_est in range(n):
                    cb = col_est * cols_por_est
                    for row in tabla.rows:
                        for c_off in range(cols_por_est):
                            c = row.cells[cb + c_off]
                            c.width = DocxCm(ancho_lbl if c_off == 0 else ancho_per)

                documento.add_paragraph("")

        documento.save(ruta)


def main() -> None:
    root = tk.Tk()
    GeneradorBoletinesApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
